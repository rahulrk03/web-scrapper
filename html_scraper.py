import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE_URL = ""

# ========== STEP 1: FETCH ALL abc IDs ==========
def fetch_abc_ids():
    url = f"{BASE_URL}"
    headers = {
        "accept": "application/json, text/plain, */*",
        "user-agent": "Mozilla/5.0",
        "referer": f"",
    }
    res = requests.get(url, headers=headers, verify=False)  # ignore internal cert check
    res.raise_for_status()
    data = res.json()
    abc_ids = [item.get("abc_id") for item in data if "abc_id" in item]
    return abc_ids


# ========== STEP 2: FETCH DETAIL PAGE ==========
def fetch_html(abc_id: int) -> str:
    url = f""
    res = requests.get(url, verify=False)
    res.raise_for_status()
    return res.text


# ========== STEP 3: PARSE DETAIL PAGE ==========
def parse_abc_detail(html: str) -> dict:
    soup = BeautifulSoup(html, "html.parser")
    data = {}

    # Title
    h3 = soup.find("h3")
    if h3:
        data["Title"] = h3.get_text(" ", strip=True)

    # Practices
    practices_block = soup.find("h4", string="Practices")
    if practices_block:
        badges = practices_block.find_all_next("span", class_="badge")
        data["Practices"] = [b.get_text(strip=True) for b in badges]

    def get_text_by_id(id_):
        el = soup.select_one(f"#{id_}")
        return el.get_text(" ", strip=True) if el else None

    # Sections
    data["Summary"] = get_text_by_id("description_summary")
    data["Primary Issue"] = get_text_by_id("description_primary_issue")
    data["Example"] = get_text_by_id("description_example")
    data["Consequence"] = get_text_by_id("consequence")
    data["Remediation Summary"] = get_text_by_id("general_remediation_summary")
    data["Primary Remediation"] = get_text_by_id("general_remediation_primary_remediation")

    # Steps to Reproduce
    steps = []
    steps_section = soup.find("h3", string=lambda s: s and "Steps to Reproduce" in s)
    if steps_section:
        container = steps_section.find_parent("div", class_="col-12 col-lg-6")
        if container:
            # Step header
            header = container.find("h4")
            if header:
                steps.append(header.get_text(" ", strip=True))
            # Ordered list items
            for step_list in container.select("ol"):
                for li in step_list.find_all("li"):
                    steps.append(li.get_text(" ", strip=True))
    data["Steps"] = steps

    # Captions
    captions = []
    for cap_list in soup.select("h5:contains('Captions') ~ ol"):
        for li in cap_list.find_all("li"):
            captions.append(li.get_text(" ", strip=True))
    data["Captions"] = captions

    # Parameters
    params = {}
    for dl in soup.select("h5:contains('Parameters') ~ dl"):
        for dt, dd in zip(dl.find_all("dt"), dl.find_all("dd")):
            params[dt.get_text(strip=True)] = dd.get_text(" ", strip=True)
    data["Parameters"] = params

    # Classifications
    classifications = {}
    for cid in ["CVSS2", "CVSS3", "CVSS4", "NIST5Likelihood", "NIST5Impact"]:
        el = soup.select_one(f"#{cid}")
        if el:
            classifications[cid] = el.get_text(" ", strip=True)
    data["Classifications"] = classifications

    # CWE
    cwes = [li.get_text(" ", strip=True) for li in soup.select("#cwes li")]
    data["CWE"] = cwes

    # CAPEC
    capecs = [li.get_text(" ", strip=True) for li in soup.select("#capecs li")]
    data["CAPEC"] = capecs

    return data


# ========== STEP 4: ADD TOC ==========
def add_table_of_contents(doc):
    """
    Inserts a Word Table of Contents field.
    After opening the file in Word, press F9 to update TOC.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # TOC for heading levels 1-3

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    doc.add_paragraph()  # add spacing after TOC


# ========== STEP 5: EXPORT TO DOCX ==========
def export_to_docx(all_data, filename="report.docx"):
    doc = Document()
    doc.add_heading("Report", 0)

    # TOC
    doc.add_heading("Table of Contents", level=1)
    add_table_of_contents(doc)
    doc.add_page_break()

    for item in all_data:
        doc.add_heading(item.get("Title", "Unknown abc"), level=1)

        if item.get("Practices"):
            doc.add_heading("Practices", level=2)
            doc.add_paragraph(", ".join(item["Practices"]))

        for section in ["Summary", "Primary Issue", "Example", "Consequence",
                        "Remediation Summary", "Primary Remediation"]:
            if item.get(section):
                doc.add_heading(section, level=2)
                doc.add_paragraph(item[section])

        # ✅ Steps formatting
        if item.get("Steps"):
            doc.add_heading("Steps to Reproduce", level=2)
            steps = item["Steps"]

            if steps:
                # First entry is header
                doc.add_paragraph(steps[0], style="Intense Quote")

                # Rest as bullets
                for step in steps[1:]:
                    if step.strip().startswith("afl") or step.strip() in ["aaaa", "r2 $main_app_binary$"]:
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(step)
                        run.font.name = "Courier New"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                        run.font.size = Pt(10)
                    else:
                        doc.add_paragraph(step, style="List Bullet")

        if item.get("Captions"):
            doc.add_heading("Captions", level=2)
            for cap in item["Captions"]:
                doc.add_paragraph(cap, style="List Bullet")

        if item.get("Parameters"):
            doc.add_heading("Parameters", level=2)
            for k, v in item["Parameters"].items():
                doc.add_paragraph(f"{k}: {v}")

        if item.get("Classifications"):
            doc.add_heading("Classifications", level=2)
            for k, v in item["Classifications"].items():
                doc.add_paragraph(f"{k}: {v}")

        if item.get("CWE"):
            doc.add_heading("CWE", level=2)
            for c in item["CWE"]:
                doc.add_paragraph(c, style="List Bullet")

        if item.get("CAPEC"):
            doc.add_heading("CAPEC", level=2)
            for c in item["CAPEC"]:
                doc.add_paragraph(c, style="List Bullet")

        doc.add_page_break()

    doc.save(filename)
    print(f"✅ Saved report to {filename}")


# ========== MAIN RUN ==========
if __name__ == "__main__":
    print("Fetching abc IDs...")
    abc_ids = sorted(fetch_abc_ids())   # ✅ order by ID ascending
    print(f"Found {len(abc_ids)} abcs (ordered)")

    all_data = []
    for cid in abc_ids:
        print(f"Fetching abc-{cid}...")
        html = fetch_html(cid)
        data = parse_abc_detail(html)
        all_data.append(data)

    export_to_docx(all_data, "abcs_report.docx")
