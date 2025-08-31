"""
Data Export Module for Web Scraping Framework

This module handles exporting scraped data to various formats including
JSON, CSV, Excel, DOCX, and plain text.
"""

import json
import csv
import os
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

# Optional imports - will handle gracefully if not available
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export disabled.")

try:
    import openpyxl
    from openpyxl import Workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl not available. Excel export disabled.")


class DataExporter:
    """
    Main data exporter class that handles multiple output formats.
    """
    
    def __init__(self, config: Optional[Dict] = None):
        """
        Initialize the data exporter.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {}
        self.output_dir = Path(self.config.get('directory', './output'))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Default filename settings
        self.base_filename = self.config.get('filename', 'scraped_data')
        self.add_timestamp = self.config.get('timestamp', True)
    
    def export(self, data: Union[Dict, List], format_type: str = None, 
               filename: str = None) -> str:
        """
        Export data to specified format.
        
        Args:
            data: Data to export
            format_type: Output format (json, csv, xlsx, docx, txt)
            filename: Custom filename (without extension)
            
        Returns:
            Path to exported file
        """
        format_type = format_type or self.config.get('format', 'json')
        filename = filename or self.base_filename
        
        # Add timestamp if enabled
        if self.add_timestamp:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{filename}_{timestamp}"
        
        # Export based on format
        if format_type == 'json':
            return self.export_json(data, filename)
        elif format_type == 'csv':
            return self.export_csv(data, filename)
        elif format_type == 'xlsx':
            return self.export_excel(data, filename)
        elif format_type == 'docx':
            return self.export_docx(data, filename)
        elif format_type == 'txt':
            return self.export_text(data, filename)
        else:
            raise ValueError(f"Unsupported export format: {format_type}")
    
    def export_json(self, data: Union[Dict, List], filename: str) -> str:
        """Export data to JSON format."""
        filepath = self.output_dir / f"{filename}.json"
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Data exported to JSON: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error exporting to JSON: {str(e)}")
            raise
    
    def export_csv(self, data: Union[Dict, List], filename: str) -> str:
        """Export data to CSV format."""
        filepath = self.output_dir / f"{filename}.csv"
        
        try:
            # Convert data to list of dictionaries if needed
            rows = self._prepare_data_for_tabular(data)
            
            if not rows:
                logger.warning("No data to export to CSV")
                return str(filepath)
            
            # Get all unique keys for headers
            headers = set()
            for row in rows:
                if isinstance(row, dict):
                    headers.update(row.keys())
            
            headers = sorted(list(headers))
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                
                for row in rows:
                    if isinstance(row, dict):
                        # Convert complex values to strings
                        clean_row = {}
                        for key in headers:
                            value = row.get(key, '')
                            if isinstance(value, (list, dict)):
                                clean_row[key] = json.dumps(value, ensure_ascii=False)
                            else:
                                clean_row[key] = str(value) if value is not None else ''
                        writer.writerow(clean_row)
            
            logger.info(f"Data exported to CSV: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error exporting to CSV: {str(e)}")
            raise
    
    def export_excel(self, data: Union[Dict, List], filename: str) -> str:
        """Export data to Excel format."""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel export")
        
        filepath = self.output_dir / f"{filename}.xlsx"
        
        try:
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = "Scraped Data"
            
            # Convert data to list of dictionaries if needed
            rows = self._prepare_data_for_tabular(data)
            
            if not rows:
                logger.warning("No data to export to Excel")
                workbook.save(filepath)
                return str(filepath)
            
            # Get headers
            headers = set()
            for row in rows:
                if isinstance(row, dict):
                    headers.update(row.keys())
            
            headers = sorted(list(headers))
            
            # Write headers
            for col, header in enumerate(headers, 1):
                worksheet.cell(row=1, column=col, value=header)
            
            # Write data
            for row_num, row in enumerate(rows, 2):
                if isinstance(row, dict):
                    for col, header in enumerate(headers, 1):
                        value = row.get(header, '')
                        if isinstance(value, (list, dict)):
                            value = json.dumps(value, ensure_ascii=False)
                        worksheet.cell(row=row_num, column=col, value=value)
            
            workbook.save(filepath)
            logger.info(f"Data exported to Excel: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            raise
    
    def export_docx(self, data: Union[Dict, List], filename: str) -> str:
        """Export data to Word document format."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
        
        filepath = self.output_dir / f"{filename}.docx"
        
        try:
            # Check if there's a template to use
            template_path = self.config.get('docx_template')
            if template_path and Path(template_path).exists():
                document = Document(template_path)
            else:
                document = Document()
            
            # Add title
            title = document.add_heading('Web Scraping Results', 0)
            
            # Add metadata
            document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            document.add_paragraph()
            
            # Process data
            if isinstance(data, dict):
                self._add_dict_to_docx(document, data)
            elif isinstance(data, list):
                for i, item in enumerate(data):
                    document.add_heading(f'Item {i + 1}', level=1)
                    if isinstance(item, dict):
                        self._add_dict_to_docx(document, item)
                    else:
                        document.add_paragraph(str(item))
                    document.add_paragraph()
            else:
                document.add_paragraph(str(data))
            
            document.save(filepath)
            logger.info(f"Data exported to DOCX: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            raise
    
    def export_text(self, data: Union[Dict, List], filename: str) -> str:
        """Export data to plain text format."""
        filepath = self.output_dir / f"{filename}.txt"
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write("Web Scraping Results\n")
                f.write("=" * 50 + "\n")
                f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                if isinstance(data, dict):
                    self._write_dict_to_text(f, data)
                elif isinstance(data, list):
                    for i, item in enumerate(data):
                        f.write(f"Item {i + 1}:\n")
                        f.write("-" * 20 + "\n")
                        if isinstance(item, dict):
                            self._write_dict_to_text(f, item)
                        else:
                            f.write(str(item) + "\n")
                        f.write("\n")
                else:
                    f.write(str(data))
            
            logger.info(f"Data exported to text: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error exporting to text: {str(e)}")
            raise
    
    def _prepare_data_for_tabular(self, data: Union[Dict, List]) -> List[Dict]:
        """Prepare data for tabular formats (CSV, Excel)."""
        if isinstance(data, dict):
            # If it's a single dictionary, wrap it in a list
            return [data]
        elif isinstance(data, list):
            # Filter for dictionaries
            return [item for item in data if isinstance(item, dict)]
        else:
            # Convert other types to dictionary
            return [{'data': str(data)}]
    
    def _add_dict_to_docx(self, document, data: Dict, level: int = 2):
        """Add dictionary data to DOCX document."""
        for key, value in data.items():
            if isinstance(value, dict):
                document.add_heading(str(key).replace('_', ' ').title(), level=level)
                self._add_dict_to_docx(document, value, level + 1)
            elif isinstance(value, list):
                document.add_heading(str(key).replace('_', ' ').title(), level=level)
                for item in value:
                    if isinstance(item, dict):
                        self._add_dict_to_docx(document, item, level + 1)
                    else:
                        p = document.add_paragraph()
                        p.add_run(f"• {str(item)}")
            else:
                p = document.add_paragraph()
                p.add_run(f"{str(key).replace('_', ' ').title()}: ").bold = True
                p.add_run(str(value))
    
    def _write_dict_to_text(self, file, data: Dict, indent: int = 0):
        """Write dictionary data to text file."""
        prefix = "  " * indent
        for key, value in data.items():
            if isinstance(value, dict):
                file.write(f"{prefix}{key}:\n")
                self._write_dict_to_text(file, value, indent + 1)
            elif isinstance(value, list):
                file.write(f"{prefix}{key}:\n")
                for item in value:
                    if isinstance(item, dict):
                        self._write_dict_to_text(file, item, indent + 1)
                    else:
                        file.write(f"{prefix}  - {str(item)}\n")
            else:
                file.write(f"{prefix}{key}: {str(value)}\n")


class BatchExporter:
    """
    Handles batch export operations for multiple data sets.
    """
    
    def __init__(self, config: Optional[Dict] = None):
        self.exporter = DataExporter(config)
        self.config = config or {}
    
    def export_multiple(self, datasets: List[Dict], base_filename: str = "batch") -> List[str]:
        """
        Export multiple datasets.
        
        Args:
            datasets: List of data dictionaries
            base_filename: Base name for files
            
        Returns:
            List of exported file paths
        """
        exported_files = []
        format_type = self.config.get('format', 'json')
        
        for i, data in enumerate(datasets):
            filename = f"{base_filename}_{i+1:03d}"
            try:
                filepath = self.exporter.export(data, format_type, filename)
                exported_files.append(filepath)
            except Exception as e:
                logger.error(f"Error exporting dataset {i+1}: {str(e)}")
        
        return exported_files
    
    def export_combined(self, datasets: List[Dict], filename: str = "combined") -> str:
        """
        Export all datasets combined into a single file.
        
        Args:
            datasets: List of data dictionaries
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        combined_data = {
            'metadata': {
                'total_datasets': len(datasets),
                'exported_at': datetime.now().isoformat()
            },
            'datasets': datasets
        }
        
        format_type = self.config.get('format', 'json')
        return self.exporter.export(combined_data, format_type, filename)


if __name__ == "__main__":
    # Example usage
    print("Data Export Module")
    print("==================")
    
    # Sample data
    sample_data = {
        'url': 'https://example.com',
        'title': 'Example Page',
        'headings': {
            'h1': ['Main Title'],
            'h2': ['Subtitle 1', 'Subtitle 2']
        },
        'links': [
            {'text': 'Link 1', 'url': 'https://example.com/page1'},
            {'text': 'Link 2', 'url': 'https://example.com/page2'}
        ],
        'text_content': 'This is sample content from the webpage.'
    }
    
    # Create exporter
    config = {
        'directory': './output',
        'filename': 'sample_export',
        'timestamp': True
    }
    
    exporter = DataExporter(config)
    
    # Export to different formats
    print("\nExporting sample data...")
    try:
        json_file = exporter.export(sample_data, 'json')
        print(f"JSON export: {json_file}")
        
        csv_file = exporter.export(sample_data, 'csv')
        print(f"CSV export: {csv_file}")
        
        txt_file = exporter.export(sample_data, 'txt')
        print(f"Text export: {txt_file}")
        
        if DOCX_AVAILABLE:
            docx_file = exporter.export(sample_data, 'docx')
            print(f"DOCX export: {docx_file}")
        
        if EXCEL_AVAILABLE:
            xlsx_file = exporter.export(sample_data, 'xlsx')
            print(f"Excel export: {xlsx_file}")
    
    except Exception as e:
        print(f"Export error: {e}")
    
    print(f"\nSupported formats:")
    print(f"- JSON: Always available")
    print(f"- CSV: Always available")
    print(f"- Text: Always available")
    print(f"- DOCX: {'Available' if DOCX_AVAILABLE else 'Requires python-docx'}")
    print(f"- Excel: {'Available' if EXCEL_AVAILABLE else 'Requires openpyxl'}")